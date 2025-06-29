from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from datetime import datetime

def create_docx(title, script, translation):
    # 문서 객체 생성
    doc = Document()

    # 제목
    doc.add_heading(title, level=1)

    # 스크립트 삽입
    doc.add_heading("스크립트", level=2)
    doc.add_paragraph(script)

    # 번역 삽입
    doc.add_heading("번역", level=2)
    doc.add_paragraph(translation)

    # 현재 작업 경로에 저장 (로컬 테스트용)
    # doc.save(f"{datetime.now().strftime('%Y%m%d')}_{title}.docx")

    # Document 객체 반환
    return doc